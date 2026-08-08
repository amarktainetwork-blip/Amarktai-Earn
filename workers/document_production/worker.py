from __future__ import annotations

import os
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches, Pt
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from workers.base import WorkRequest, WorkResult, Worker


def _content(inputs: dict) -> tuple[str, str]:
    source = inputs.get("source")
    text = Path(str(source)).read_text(encoding="utf-8", errors="replace") if source else str(inputs.get("content") or "")
    title = str(inputs.get("title") or "Professional Document").strip()[:200]
    text = text.replace("\x00", "").strip()
    if not text:
        raise ValueError("document content is required")
    maximum = max(1000, int(os.getenv("DOCUMENT_PRODUCTION_MAX_CHARS", "500000")))
    if len(text) > maximum:
        raise ValueError("DOCUMENT_PRODUCTION_CONTENT_LIMIT")
    return title, text


class DocumentProductionWorker(Worker):
    worker_class = "document_production"

    def execute(self, request: WorkRequest) -> WorkResult:
        try:
            operation = str(request.inputs.get("operation") or "")
            if operation not in {"docx_create", "pdf_create"}:
                return WorkResult(ok=False, error="unsupported document-production operation")
            title, text = _content(request.inputs)
            request.workspace.mkdir(parents=True, exist_ok=True)
            if operation == "docx_create":
                target = request.workspace / "document.docx"
                document = Document(); section = document.sections[0]
                section.top_margin = section.bottom_margin = Inches(0.75)
                document.add_heading(title, 0)
                for block in text.split("\n\n"):
                    block = block.strip()
                    if not block: continue
                    if block.startswith("#"):
                        document.add_heading(block.lstrip("# ")[:250], level=min(3, len(block) - len(block.lstrip("#"))))
                    else:
                        paragraph = document.add_paragraph(block)
                        paragraph.style.font.size = Pt(10.5)
                document.core_properties.title = title; document.save(target)
                reopened = Document(target)
                evidence = {"paragraph_count": len(reopened.paragraphs), "reopened": True, "format": "docx"}
            else:
                target = request.workspace / "document.pdf"
                styles = getSampleStyleSheet()
                story = [Paragraph(title, styles["Title"]), Spacer(1, 0.2 * inch)]
                for block in text.split("\n\n"):
                    cleaned = block.strip()
                    if cleaned: story.extend([Paragraph(escape(cleaned).replace("\n", "<br/>"), styles["BodyText"]), Spacer(1, 0.12 * inch)])
                SimpleDocTemplate(str(target), pagesize=A4, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch, title=title).build(story)
                reader = PdfReader(str(target))
                evidence = {"page_count": len(reader.pages), "reopened": True, "format": "pdf"}
            evidence.update({"operation": operation, "source_chars": len(text), "title": title})
            return WorkResult(ok=True, artifacts=[target], evidence=evidence)
        except (OSError, KeyError, TypeError, ValueError) as exc:
            return WorkResult(ok=False, error=str(exc))
